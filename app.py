import streamlit as st
import json
import uuid
import os
import pathlib
import streamlit.components.v1 as components
from backend.gemini_service import call_gemini,call_gemini_to_export_sql,call_gemini_to_get_documentation_for_model
import datetime
from docx import Document
from io import BytesIO

# Set up the page configuration
st.set_page_config(
    page_title="DAXI",
    layout="wide",
    initial_sidebar_state="expanded"
)
st.markdown("""
    <style>
    .stButton > button
     {
        background-color: #4c4f52;
        color: white;
    }
    .stButton > button:hover {
        background-color: #F79B72;
    }
    </style>
""", unsafe_allow_html=True)

# Define the directory for saving projects
PROJECTS_DIR = pathlib.Path("projects")
PROJECTS_DIR.mkdir(exist_ok=True)

REACT_PROJECTS_DIR = pathlib.Path("react_flow_component/src/projects")
REACT_PROJECTS_DIR.mkdir(exist_ok=True)

# Helper function to get saved projects
def get_saved_projects():
    """Scans the projects directory and returns a list of project names."""
    projects = [f.stem for f in PROJECTS_DIR.glob("*.json")]
    return sorted(projects)

# --- Initialize all session state variables at the top of the script ---
if 'current_project' not in st.session_state:
    st.session_state.current_project = None
if 'project_history' not in st.session_state:
    st.session_state.project_history = []
if 'history_index' not in st.session_state:
    st.session_state.history_index = -1
if 'new_project_name' not in st.session_state:
    st.session_state.new_project_name = ""
if 'projects_list' not in st.session_state:
    st.session_state.projects_list = []
if 'selected_project' not in st.session_state:
    st.session_state.selected_project = None
if 'json_string' not in st.session_state:
    st.session_state.json_string = ""


# --- Check URL Query Parameters for persistence after initialization ---
query_params = st.query_params
initial_project_name = query_params.get("project", None)

if initial_project_name and st.session_state.current_project is None:
    file_path = PROJECTS_DIR / f"{initial_project_name.replace(' ', '_').lower()}.json"
    if file_path.exists():
        try:
            with open(file_path, "r") as f:
                loaded_schema = json.load(f)
            st.session_state.current_project = {
                "id": str(uuid.uuid4()),
                "name": initial_project_name,
                "data": loaded_schema
            }
            st.session_state.project_history = []
            st.session_state.history_index = -1
            st.session_state.json_string = json.dumps(st.session_state.current_project['data'], indent=4)
            # st.success(f"Project '{initial_project_name}' loaded from URL.")
        except (FileNotFoundError, json.JSONDecodeError) as e:
            st.error(f"Error loading project from URL: {e}")
    else:
        st.error(f"Project '{initial_project_name}' not found.")
        st.experimental_set_query_params() # Clear the URL param to avoid an infinite loop


def save_state():
    """Saves the current project state to the in-memory history for undo/redo."""
    if st.session_state.current_project is not None:
        update_curr_file()
        # Clear any "redo" states if a new action is performed
        if st.session_state.history_index + 1 < len(st.session_state.project_history):
            st.session_state.project_history = st.session_state.project_history[:st.session_state.history_index + 1]

        # Add the current state to the history
        st.session_state.project_history.append(json.dumps(st.session_state.current_project, indent=4))
        st.session_state.history_index = len(st.session_state.project_history) - 1

def save_to_file():
    """Saves the current project to a JSON file in the projects directory."""
    if st.session_state.current_project is not None:
        update_curr_file()
        project_name = st.session_state.current_project.get("name", "untitled_project")
        file_path = PROJECTS_DIR / f"{project_name.replace(' ', '_').lower()}.json"
        
        # Save only the 'data' part of the project object
        with open(file_path, "w") as f:
            json.dump(st.session_state.current_project['data'], f, indent=4)
        st.session_state.projects_list = get_saved_projects()
        st.success(f"Project '{project_name}' saved to disk.")

def update_curr_file():
    """Saves the current project to a JSON file in the projects directory."""
    if st.session_state.current_project is not None:
        project_name = "curr_model"
        file_path = REACT_PROJECTS_DIR / f"{project_name.replace(' ', '_').lower()}.json"
        
        # Save only the 'data' part of the project object
        with open(file_path, "w") as f:
            json.dump(st.session_state.current_project['data'], f, indent=4)
                


def undo():
    """Moves back one step in the project history."""
    if st.session_state.history_index > 0:
        update_curr_file()
        st.session_state.history_index -= 1
        st.session_state.current_project = json.loads(st.session_state.project_history[st.session_state.history_index])
        st.session_state.json_string = json.dumps(st.session_state.current_project['data'], indent=4)
    else:
        st.warning("Cannot undo further.")

def redo():
    """Moves forward one step in the project history."""
    if st.session_state.history_index < len(st.session_state.project_history) - 1:
        update_curr_file()
        st.session_state.history_index += 1
        st.session_state.current_project = json.loads(st.session_state.project_history[st.session_state.history_index])
        st.session_state.json_string = json.dumps(st.session_state.current_project['data'], indent=4)
    else:
        st.warning("Cannot redo further.")

def create_new_project():
    """Creates a new, empty project."""
    project_id = str(uuid.uuid4())
    project_name = st.session_state.new_project_name if st.session_state.new_project_name else "New Project"
    
    st.session_state.current_project = {
        "id": project_id,
        "name": project_name,
        "data": {
            "schema": {
                "entities": [],
                "relationships": []
            }
        }
    }
    
    # Update the URL to reflect the new project name
    st.query_params["project"] = project_name
    
    # Reset history for the new project
    st.session_state.project_history = []
    st.session_state.history_index = -1
    save_state()
    st.session_state.new_project_name = "" # Clear the input field
    st.session_state.json_string = json.dumps(st.session_state.current_project['data'], indent=4)
    # st.success(f"Project '{project_name}' created successfully!")


def load_project():
    """Loads a project from the projects folder based on the selected name."""
    if st.session_state.selected_project:
        file_path = PROJECTS_DIR / f"{st.session_state.selected_project}.json"
        try:
            with open(file_path, "r") as f:
                loaded_schema = json.load(f)

            # Recreate the full project object
            project_id = str(uuid.uuid4())
            project_name = st.session_state.selected_project
            st.session_state.current_project = {
                "id": project_id,
                "name": project_name,
                "data": loaded_schema
            }

            # Update the URL to reflect the loaded project name
            st.query_params["project"] = project_name

            # Reset history for the new project
            st.session_state.project_history = []
            st.session_state.history_index = -1
            save_state()
            st.session_state.json_string = json.dumps(st.session_state.current_project['data'], indent=4)
            # st.success(f"Project '{st.session_state.selected_project}' loaded successfully!")
        except (FileNotFoundError, json.JSONDecodeError) as e:
            st.error(f"Error loading project: {e}")
    else:
        st.warning("Please select a project to load.")

def handle_json_update():
    """Handles parsing and updating the state when the text area changes."""
    try:
        updated_data = json.loads(st.session_state.json_string)
        st.session_state.current_project['data'] = updated_data
        save_state()
    except json.JSONDecodeError:
        st.error("Invalid JSON format. Please correct the syntax.")

def rename_project():
    """Renames the current project and its corresponding file."""
    if 'rename_project_name' in st.session_state and st.session_state.rename_project_name:
        old_name = st.session_state.current_project.get("name")
        new_name = st.session_state.rename_project_name
        
        if old_name and old_name != new_name:
            old_file_path = PROJECTS_DIR / f"{old_name.replace(' ', '_').lower()}.json"
            new_file_path = PROJECTS_DIR / f"{new_name.replace(' ', '_').lower()}.json"
            
            try:
                if old_file_path.exists():
                    os.rename(old_file_path, new_file_path)
                
                st.session_state.current_project['name'] = new_name
                st.query_params["project"] = new_name # Update URL
                save_state()
                st.session_state.projects_list = get_saved_projects()
                st.success(f"Project renamed to '{new_name}'!")
            except OSError as e:
                st.error(f"Error renaming project file: {e}")
        else:
            st.warning("Please enter a new name to rename the project.")

# --- UI Layout ---

st.title("DAXI")

# Sidebar for project management
with st.sidebar:
    
    st.subheader("Create New Project")
    st.text_input("Project Name", key="new_project_name", help="Leave blank for a default name.")
    if st.button("Create", on_click=create_new_project, use_container_width=True):
        pass

    st.subheader("Load Saved Project")
    st.session_state.projects_list = get_saved_projects()
    st.selectbox(
        "Select a project",
        options=st.session_state.projects_list,
        key="selected_project",
        index=None
    )
    if st.button("Load", on_click=load_project, use_container_width=True, disabled=not st.session_state.selected_project):
        pass

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Undo", use_container_width=True, disabled=st.session_state.history_index <= 0):
            undo()
    with col2:
        if st.button("Redo", use_container_width=True, disabled=st.session_state.history_index >= len(st.session_state.project_history) - 1):
            redo()

# Main content area
if st.session_state.current_project is not None:
    # Add rename functionality
    with st.container():
        col_rename1, col_rename2,col_dropdown, col_export, col_documentation = st.columns([2, 1, 1, 1, 1])
        with col_rename1:
            st.text_input(
                "Rename Project Name", 
                value=st.session_state.current_project.get("name", ""),
                key="rename_project_name", 
                help="Enter a new name for the project."
            )
        with col_rename2:
            st.markdown("<br>", unsafe_allow_html=True)
            st.button("Rename", on_click=rename_project, use_container_width=True)
        
        with col_dropdown:
            export_format = st.selectbox(
                "Export To:",
                options=["Oracle", "SQL Server", "Postgre", "MongoDB", "DB2"], # Changed default option text
                index=0,
                # label_visibility="collapsed"
            )
            
        with col_export:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Get DDL",  use_container_width=True):
                if export_format == "Select DB": # Check for the specific default dropdown text
                    st.warning("Please select a database type for export.")
                else:
                    # Use st.spinner for better UX during the API call
                    with st.spinner(f"Generating DDL for {export_format}... Please wait."):
                        payload = {
                            "db_type": export_format,
                        }
                        print("Payload for DDL generation:")
                        print(payload)

                        try:
                            # Call your specific API for SQL export
                            # Ensure `call_api_to_export_sql` returns the raw SQL string
                            sql_response_content = call_gemini_to_export_sql(
                                db_type=export_format,
                                currFile=st.session_state.current_project['data']
                            )
                            download_filename = f"{export_format.lower()}_ddl_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.sql"
                            mime_type = "application/sql"
                            # Convert content to bytes
                            st.download_button(
                                label="Download the File",
                                data=sql_response_content,
                                file_name=download_filename,
                                mime=mime_type
                            )

                        except Exception as e:
                            st.error(f"An unexpected error occurred during DDL generation: {e}")

        with col_documentation:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("Get Doc",  use_container_width=True):
                with st.spinner(f"Generating Documentation... Please wait."):
                    try:
                        documentation_content = call_gemini_to_get_documentation_for_model(
                            currFile=st.session_state.current_project['data']
                        )
                        
                        doc = Document()
                        doc.add_heading('Generated Documentation', level=1)
                        for line in documentation_content.strip().split('\n'):
                            doc.add_paragraph(line)
                            
                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)

                        download_filename = f"documentation_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        # Convert content to bytes
                        st.download_button(
                            label="Download the File",
                            data=buffer,
                            file_name=download_filename,
                            mime=mime_type
                        )

                    except Exception as e:
                        st.error(f"An unexpected error occurred during DDL generation: {e}")

    st.markdown("---")
    erd_col,json_col = st.columns([4,1])
    
    with json_col:
        # Use st.text_area to allow direct editing of the JSON string
        if st.button("Save", use_container_width=True):
            save_to_file()
        # st.text_area(
        #     "",
        #     value=st.session_state.json_string,
        #     height=450,
        #     key="json_string",
        #     on_change=handle_json_update
        # )
    
    with erd_col:
        # NOTE: The ERD IFrame will only work if you have a separate server
        # running at http://localhost:3000/ to serve the ERD visualization.
        # This part of the code is kept from your original request.
        iframe_base_url = f"http://localhost:3000/"

        if 'iframe_current_url' not in st.session_state:
            st.session_state.iframe_current_url = iframe_base_url
        
        components.iframe(st.session_state.iframe_current_url, height=500, scrolling=True)
    
    

    Mtab1 = st.tabs(["📝 Prompt"])[0]

    with Mtab1:
        # --- Tab 1: Prompt ---
        user_prompt = st.text_area("Type your description here...", height=150)
        tab2, tab3 = st.columns(2)
        # --- Tab 2: Upload BRD ---
        with tab2:
            brd_files = st.file_uploader("Upload Reference files:", type=["pdf", "docx"], accept_multiple_files=True)

        # --- Tab 3: Upload Schema ---
        with tab3:
            schema_files = st.file_uploader("Upload Legacy schema files:", type=["csv", "xlsx", "xls", "sql", "json"], accept_multiple_files=True)

    # --- Generate Button centered ---
    st.markdown('<div class="center-button">', unsafe_allow_html=True)
    generate_clicked = st.button("🛠️ Generate")
    st.markdown('</div>', unsafe_allow_html=True)

    if generate_clicked:
        print("generate clicked")
        if user_prompt or brd_files or schema_files:
            payload = {
                "prompt": user_prompt,
                "curr_model": st.session_state.current_project['data']
            }
            # Prepare the files to be sent
            files_to_send = []
            if brd_files:
                for f in brd_files:
                    files_to_send.append(
                        ("brd_files", (f.name, f.getvalue(), f.type))
                    )
            if schema_files:
                for f in schema_files:
                    files_to_send.append(
                        ("schema_files", (f.name, f.getvalue(), f.type))
                    )
            
            print(files_to_send)
            print(payload)
            response = call_gemini(prompt=user_prompt,curr_model=st.session_state.current_project['data'],
                                   files=brd_files + schema_files # Pass both sets of files
                                   )
            print("output from generate")
            print(response)
            if response is not None:
                st.session_state.json_string = json.dumps(response)
                handle_json_update()
            st.write(response)
            st.success("Generating ER Diagram... Please wait.")

        else:
            st.warning("Please enter a prompt or upload files.")
        st.session_state.iframe_current_url = iframe_base_url
        st.cache_data.clear()
        st.rerun()
    
else:
    st.info("No project loaded. Please create a new project or load a saved one from the sidebar.")



